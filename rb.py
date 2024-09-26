import docx
document = docx.Document()

pi = [{"Name":"Nyamtulla","Phone":"9847346132"}]
p = document.add_heading(level=1)
run= p.add_run("Nyamtulla")
run.bold = True
run.underline = True
edu = [{"uni": "ku", "gpa": "3.6"},{"uni": "ku","gpa":"3.5"}]
parts = [pi,edu]
for part in parts:
    for p in part:
        for tag in p:
            document.add_paragraph(tag+": "+p[tag], style='List Bullet')


document.save('resume1.docx')

