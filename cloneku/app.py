from flask import Flask, render_template, request, send_file
import os
import openai
import docx
from bs4 import BeautifulSoup
import PyPDF2
import io
import textract


openai.organization = "org-gG6K1j8fah4HfsIk7JZEGNBO"
#openai.api_key = "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"
#openai.api_key = "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"
openai.api_key = "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"
app = Flask(__name__)

global resume_text
global cover_text


@app.route('/')
def home():
    return render_template('index.html')


@app.route('/out', methods=['GET', 'POST'])
def out():
    # Get user input from form
    user_string = request.form['user_string']
    user_file = request.files['user_file']

    # Create a PDF reader object
    reader = PyPDF2.PdfReader(user_file)

    # Get the number of pages in the PDF file
    num_pages = len(reader.pages)

    # Loop through each page and extract the text
    for i in range(num_pages):
        # Get the page object
        page = reader.pages[i]
        # Extract the text from the page
        # text = textract.process(, method='pdfminer')
        text = page.extract_text()
        text = "".join(text.split(" "))
    # print(text)

    def resume_ai(user_string, text):
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": "re write my complete resume {}, make sure you include all my skills  and tailor it according to the job description: {}".format(text, user_string)
                 }
            ],)
        resume = response['choices'][0]['message']['content']
        return resume

    def coverletter_ai(user_string, text):
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": "write a cover letter for my resume {}, make sure you tailor it according to the job description: {} ".format(text, user_string)
                 }
            ],)
        coverletter = response['choices'][0]['message']['content']
        # soup = BeautifulSoup(resume)
        # clean_resume = soup.get_text()
        return coverletter

    global resume_text
    resume_text = resume_ai(user_string, text)
    global cover_text
    cover_text = coverletter_ai(user_string, text)
    return render_template('newresume.html')


@app.route('/oresume', methods=['GET', 'POST'])
def oresume():
    r_document = docx.Document()
    r_document.add_paragraph(resume_text)
    r_document.save('resume.docx')
    return send_file("resume.docx", mimetype=None, as_attachment=True, download_name="resume")
    return "sf"


@app.route('/ocover', methods=['GET', 'POST'])
def ocover():
    r_document = docx.Document()
    r_document.add_paragraph(cover_text)
    r_document.save('cover.docx')
    return send_file("cover.docx", mimetype=None, as_attachment=True, download_name="cover")
