from flask import Flask, render_template, request, send_file
import os
import openai
import docx
from bs4 import BeautifulSoup
import PyPDF2
import io
import io

openai.organization = "org-gG6K1j8fah4HfsIk7JZEGNBO"
openai.api_key = "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"
app = Flask(__name__)


@app.route('/')
def home():
    return render_template('index.html')


@app.route('/out', methods=['GET', 'POST'])
def out():
    # Get user input from form
    user_string = request.form['user_string']
    user_file = request.files['user_file']
    # print(io.TextIOWrapper(user_file.read()))
    # file_path = os.path.join(app.root_path, 'uploads', user_file.filename)
    # user_file.save(file_path)
    # file_path = "".join(file_path.split())
    # with open(io.TextIOWrapper(user_file.read()), 'rb') as f:
    # Create a PDF reader object
    reader = PyPDF2.PdfReader(user_file)

    # Get the number of pages in the PDF file
    num_pages = len(reader.pages)

    # Loop through each page and extract the text
    for i in range(num_pages):
        # Get the page object
        page = reader.pages[i]
        # Extract the text from the page
        text = page.extract_text()
    soup = BeautifulSoup(text, features="lxml")
    covertxt = soup.get_text()
# job="https://careers.microsoft.com/students/us/en/job/1540490/Software-Engineer-Full-Time-Opportunities-Microsoft-Leap"
    # cl_response = openai.ChatCompletion.create(
    #     model="gpt-3.5-turbo",
    #     messages=[
    #         {"role": "system",
    #          "content": "write a cover letter for my resume {}, make sure you tailor it according to the job description: {} ".format(covertxt, user_string)
    #          }
    #     ],
    #     # max_tokens=500
    # )
    # coverletter = cl_response['choices'][0]['message']['content']

    re_response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system",
             "content": "re write my complete resume {}, make sure you include all my skills  and tailor it according to the job description: {}".format(covertxt, user_string)
             }
        ],
        # max_tokens=500
    )
    resume = re_response['choices'][0]['message']['content']
    # print("Hi")
    # print(coverletter)
    # soup = BeautifulSoup(coverletter, features="lxml")
    # covertxt = soup.get_text()
    # cl_document = docx.Document()
    # cl_document.add_paragraph(coverletter)
    # cl_document.save('Coetter.docx')
    r_document = docx.Document()
    r_document.add_paragraph(resume)
    r_document.save('resume.docx')
    # , send_file("resume.docx", mimetype=None, as_attachment=True, download_name="download"))
    return send_file("resume.docx", mimetype=None, as_attachment=True, download_name="resume")


if __name__ == '__main__':
    app.run(debug=True)
