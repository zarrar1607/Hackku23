import docx
from docx import Document
# Create a new document
document = docx.Document()

# Add a heading to the document
document.add_heading(
    'Abdul Baseer. Mohammed - Software Engineer Profile', level=1)

# Add personal information to the document
document.add_paragraph('Lawrence, Kansas KS')
document.add_paragraph('(785)423-8596, abdulbaseer@ku.edu')
document.add_paragraph('Linkedin: [insert linkedin url]')
document.add_paragraph('Github: [insert github url]')

# Add education details to the document
document.add_heading('Education', level=2)
document.add_paragraph(
    'Master of Science - Computer Science – University of Kansas, Lawrence, KS, (December 2023) Ongoing')
document.add_paragraph('GPA: 3.9/4.0')
document.add_paragraph(
    'Selected Coursework:[Information Retrieval, Analysis of Algorithms,Distributed Applications,Machine Learning(ML)]')
document.add_paragraph(
    'Bachelor of Engineering – Computer Science – Osmania University, Hyderabad, India, (May 2021)')
document.add_paragraph('GPA: 7.0/10.0')
document.add_paragraph(
    'Selected Coursework: [Data Structures,object oriented programing(oop), Software Engineering, Web Programing, DBMS]')

# Add skills details to the document
document.add_heading('Skills', level=2)
document.add_paragraph(
    'Programming Languages: Python, C++, Java, Node.js, Express, React.js, JavaScript, HTML, CSS,Flask, Pytorch,Django.')
document.add_paragraph(
    'Database: MySQL-relational database -Rdbms,Oracle SQL ,AWS RDS, MongoDB-NOSQL,AWS DynamoDB,PL/SQL.')
document.add_paragraph(
    'Tools and Technologies:AWS Cloud ,GIT, Github, Jupyter Notebook,Tableau, Docker, NPM, Postman, SDLC CI/CD.')

# Add professional experience details to the document
document.add_heading('Professional Experience', level=2)

# Add first job details to the document
document.add_heading(
    'Software Developer, University of Kansas, Lawrence, KS', level=3)
document.add_paragraph('January 2023 – present')
document.add_paragraph(
    'Architect and implement full stack web features in JavaScript MERN Stack (MySQL) following Agile methodology.')
document.add_paragraph(
    'Designing Relations in MySQL Database (RDBMS) and migrating existing databases to AWS RDS.')
document.add_paragraph(
    'Implement AWS Lambda functions to connect, store and map data in S3 from MySQL RDS.')
document.add_paragraph(
    'Build well-designed, reusable UI components using React.js, HTML, CSS.')
document.add_paragraph(
    'Re-engineering and optimizing backend using Node.js, Express.js to build RESTful API.')

# Add second job details to the document
document.add_heading(
    'Teaching Assistant, University of Kansas, Lawrence, KS', level=3)
document.add_paragraph('August 2022 – December 2022')
document.add_paragraph('Leverage in-depth proficiency in JavaScript to assist with EECS 368 (Programming Paradigms) while providing practical insight into innumerable aspects of modern programming languages that led to 70% better outcomes.')
document.add_paragraph(
    'Perform all aspects of instructor duties, including explaining programming assignments, reviewing Code, and conducting help sessions with 100% accuracy.')

document.save('resume.docx')
